from docx import Document

def parse_basics(doc):
    table = doc.tables[0]

    def clean_row(row):
        """Remove consecutive duplicates and empty cells."""
        texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        cleaned = []
        for t in texts:
            if not cleaned or t != cleaned[-1]:
                cleaned.append(t)
        return cleaned

    # Build key-value dictionary from table
    table_dict = {}
    for row in table.rows:
        cleaned = clean_row(row)
        if len(cleaned) >= 2:
            key, value = cleaned[0], cleaned[1]
            table_dict[key] = value

    # Map table_dict into structured JSON
    basics_json = {
        "coursename": table_dict.get("Course Unit Title", ""),
        "coursecode": table_dict.get("Course Unit Code", ""),
        "coursetype": "Compulsory" if "Compulsory" in table_dict.get("Type of Course Unit", "") else table_dict.get("Type of Course Unit", ""),
        "prerequisite": table_dict.get("Prerequisities and co-requisities", ""),
        "theoretical": table_dict.get("Theoretical (hour/week)", ""),
        "practice": table_dict.get("Practice (hour/week)", ""),
        "labcre": table_dict.get("Laboratory (hour/week)", ""),
        "yearofstudy": table_dict.get("Year of Study", ""),
        "semdel": table_dict.get("Semester when the course unit is delivered", ""),
        "level": table_dict.get("Level of Course Unit", ""),
        "eligdep": [
            dep.strip() for dep in 
            table_dict.get("Type of Course Unit", "")
            .replace("for all departments","All departments")
            .split(",")
        ],
        "lang": table_dict.get("Language of Instruction", ""),
        "mode": [m.strip() for m in table_dict.get("Mode of Delivery", "").split(",")],
        "recom": table_dict.get("Recommended Optional Programme Components", "")
    }

    return basics_json














# from docx import Document

# def parse_basics(doc):
#     table = doc.tables[0]
#     table_dict = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in table.rows}

#     basics_json = {
#         "coursename": table_dict.get("Course Unit Title", ""),
#         "coursecode": table_dict.get("Course Unit Code", ""),
#         "coursetype": "Compulsory" if "Compulsory" in table_dict.get("Type of Course Unit", "") else table_dict.get("Type of Course Unit", ""),
#         "prerequisite": table_dict.get("Prerequisites and co-requisites", ""),
#         "theoretical": table_dict.get("Theoretical (hour/week)", ""),
#         "practice": table_dict.get("Practice (hour/week)", ""),
#         "labcre": table_dict.get("Laboratory (hour/week)", ""),
#         "yearofstudy": table_dict.get("Year of Study", ""),
#         "semdel": table_dict.get("Semester when the course unit is delivered", ""),
#         "level": table_dict.get("Level of Course Unit", ""),
#         "eligdep": [dep.strip() for dep in table_dict.get("Type of Course Unit", "").replace("for all departments","All departments").split(",")],
#         "lang": table_dict.get("Language of Instruction", ""),
#         "mode": [m.strip() for m in table_dict.get("Mode of Delivery", "").split(",")],
#         "recom": table_dict.get("Recommended Optional Programme Components", "")
#     }

#     return basics_json
















# from docx import Document

# def parse_basics(doc):
#     table = doc.tables[0]

#     # Normalize each row by removing consecutive duplicate text
#     def clean_row_text(row):
#         texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
#         cleaned = []
#         for t in texts:
#             if not cleaned or t != cleaned[-1]:
#                 cleaned.append(t)
#         return cleaned

#     # Build dictionary safely
#     table_dict = {}
#     for row in table.rows:
#         cleaned = clean_row_text(row)
#         if len(cleaned) >= 2:
#             key, value = cleaned[0], cleaned[1]
#             table_dict[key] = value

#     basics_json = {
#         "coursename": table_dict.get("Course Unit Title", ""),
#         "coursecode": table_dict.get("Course Unit Code", ""),
#         "coursetype": "Compulsory" if "Compulsory" in table_dict.get("Type of Course Unit", "") else table_dict.get("Type of Course Unit", ""),
#         "prerequisite": table_dict.get("Prerequisites and co-requisites", ""),
#         "theoretical": table_dict.get("Theoretical (hour/week)", ""),
#         "practice": table_dict.get("Practice (hour/week)", ""),
#         "labcre": table_dict.get("Laboratory (hour/week)", ""),
#         "yearofstudy": table_dict.get("Year of Study", ""),
#         "semdel": table_dict.get("Semester when the course unit is delivered", ""),
#         "level": table_dict.get("Level of Course Unit", ""),
#         "eligdep": [dep.strip() for dep in table_dict.get("Type of Course Unit", "").replace("for all departments","All departments").split(",")],
#         "lang": table_dict.get("Language of Instruction", ""),
#         "mode": [m.strip() for m in table_dict.get("Mode of Delivery", "").split(",")],
#         "recom": table_dict.get("Recommended Optional Programme Components", "")
#     }

#     return basics_json
